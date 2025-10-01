#!/usr/bin/env python3
"""
Worker Performance Testing Script

This script tests PDF converter performance with different loads.
It only monitors and reports - does NOT change service configuration.

Usage Examples:
  python test_workers.py                    # Test localhost:8000 (default)
  python test_workers.py --port 8080        # Test localhost:8080
  python test_workers.py --host 192.168.1.100 --port 8000
  python test_workers.py --url http://remote-server:8000
  python test_workers.py --light            # Run only light tests
  python test_workers.py --heavy            # Run only heavy tests

What it does:
- Sends test conversion requests
- Measures response times and success rates
- Monitors CPU, memory, worker utilization
- Provides performance recommendations
- Does NOT modify service settings
"""
import requests
import time
import threading
import statistics
from concurrent.futures import ThreadPoolExecutor
import json
import os
import argparse

# Default configuration - can be overridden
DEFAULT_HOST = "localhost"
DEFAULT_PORT = 8000
BASE_URL = f"http://{DEFAULT_HOST}:{DEFAULT_PORT}"

def create_test_docx():
    """Create a simple test DOCX file"""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Performance Test Document', 0)
        doc.add_paragraph('This is a test document for worker performance testing.')
        doc.add_paragraph('Generated automatically for load testing.')
        
        # Add some content to make it more realistic
        for i in range(5):
            doc.add_paragraph(f'Paragraph {i+1}: Lorem ipsum dolor sit amet, consectetur adipiscing elit.')
        
        doc.save('test_worker_doc.docx')
        return True
    except ImportError:
        print("❌ python-docx not installed. Creating dummy file...")
        with open('test_worker_doc.docx', 'wb') as f:
            f.write(b'PK\x03\x04' + b'\x00' * 100)
        return False

def send_conversion_request():
    """Send a single conversion request and measure time"""
    start_time = time.time()
    
    try:
        with open('test_worker_doc.docx', 'rb') as f:
            files = {'file': f}
            data = {
                'nomor_urut': f'PERF_TEST_{int(time.time() * 1000)}',
                'target_url': 'http://test.example.com/callback'
            }
            
            response = requests.post(f"{BASE_URL}/convertDua", files=files, data=data, timeout=30)
            
            if response.status_code == 200:
                result = response.json()
                conversion_id = result.get('conversion_id')
                
                # Wait for completion
                max_wait = 120  # 2 minutes max
                waited = 0
                
                while waited < max_wait:
                    status_response = requests.get(f"{BASE_URL}/status/{conversion_id}", timeout=10)
                    if status_response.status_code == 200:
                        status = status_response.json()
                        if status['status'] == 'completed':
                            total_time = time.time() - start_time
                            # Cleanup
                            requests.delete(f"{BASE_URL}/cleanup/{conversion_id}")
                            return {'success': True, 'time': total_time, 'conversion_id': conversion_id}
                        elif status['status'] == 'failed':
                            return {'success': False, 'error': 'Conversion failed', 'time': time.time() - start_time}
                    
                    time.sleep(2)
                    waited += 2
                
                return {'success': False, 'error': 'Timeout', 'time': time.time() - start_time}
            else:
                return {'success': False, 'error': f'HTTP {response.status_code}', 'time': time.time() - start_time}
                
    except Exception as e:
        return {'success': False, 'error': str(e), 'time': time.time() - start_time}

def get_system_stats():
    """Get current system statistics"""
    try:
        health_response = requests.get(f"{BASE_URL}/health", timeout=10)
        if health_response.status_code == 200:
            return health_response.json()
    except:
        pass
    return None

def run_load_test(concurrent_requests=5, total_requests=20):
    """Run load test with specified parameters"""
    print(f"\n🚀 Running load test: {concurrent_requests} concurrent, {total_requests} total requests")
    
    # Get initial system stats
    initial_stats = get_system_stats()
    if initial_stats:
        print(f"   Initial CPU: {initial_stats['system_resources']['cpu_percent']:.1f}%")
        print(f"   Initial Memory: {initial_stats['system_resources']['memory_percent']:.1f}%")
        print(f"   Max Workers: {initial_stats['max_workers']}")
    
    results = []
    start_time = time.time()
    
    # Use ThreadPoolExecutor to send concurrent requests
    with ThreadPoolExecutor(max_workers=concurrent_requests) as executor:
        # Submit all requests
        futures = [executor.submit(send_conversion_request) for _ in range(total_requests)]
        
        # Collect results
        for i, future in enumerate(futures):
            result = future.result()
            results.append(result)
            
            status = "✅" if result['success'] else "❌"
            print(f"   Request {i+1:2d}: {status} {result['time']:.2f}s")
    
    total_time = time.time() - start_time
    
    # Calculate statistics
    successful = [r for r in results if r['success']]
    failed = [r for r in results if not r['success']]
    
    if successful:
        times = [r['time'] for r in successful]
        avg_time = statistics.mean(times)
        median_time = statistics.median(times)
        min_time = min(times)
        max_time = max(times)
    else:
        avg_time = median_time = min_time = max_time = 0
    
    # Get final system stats
    final_stats = get_system_stats()
    
    print(f"\n📊 Results Summary:")
    print(f"   Total time: {total_time:.2f}s")
    print(f"   Successful: {len(successful)}/{total_requests} ({len(successful)/total_requests*100:.1f}%)")
    print(f"   Failed: {len(failed)}")
    print(f"   Average time: {avg_time:.2f}s")
    print(f"   Median time: {median_time:.2f}s")
    print(f"   Min time: {min_time:.2f}s")
    print(f"   Max time: {max_time:.2f}s")
    print(f"   Throughput: {len(successful)/total_time:.2f} conversions/sec")
    
    if final_stats:
        print(f"   Final CPU: {final_stats['system_resources']['cpu_percent']:.1f}%")
        print(f"   Final Memory: {final_stats['system_resources']['memory_percent']:.1f}%")
        print(f"   Worker Utilization: {final_stats['worker_utilization']}")
    
    if failed:
        print(f"\n❌ Failed requests:")
        for i, fail in enumerate(failed[:5]):  # Show first 5 failures
            print(f"   {i+1}. {fail['error']}")
    
    return {
        'total_requests': total_requests,
        'successful': len(successful),
        'failed': len(failed),
        'success_rate': len(successful)/total_requests*100,
        'avg_time': avg_time,
        'median_time': median_time,
        'throughput': len(successful)/total_time,
        'total_time': total_time,
        'initial_stats': initial_stats,
        'final_stats': final_stats
    }

def main():
    """Run comprehensive worker performance tests"""
    global BASE_URL
    
    # Parse command line arguments
    parser = argparse.ArgumentParser(description='PDF Converter Worker Performance Testing')
    parser.add_argument('--host', default=DEFAULT_HOST, help='Service host (default: localhost)')
    parser.add_argument('--port', type=int, default=DEFAULT_PORT, help='Service port (default: 8000)')
    parser.add_argument('--url', help='Full service URL (overrides host/port)')
    parser.add_argument('--light', action='store_true', help='Run only light tests')
    parser.add_argument('--heavy', action='store_true', help='Run only heavy tests')
    
    args = parser.parse_args()
    
    # Set BASE_URL based on arguments
    if args.url:
        BASE_URL = args.url.rstrip('/')
    else:
        BASE_URL = f"http://{args.host}:{args.port}"
    
    print("🔧 Worker Performance Testing Tool")
    print("=" * 50)
    print(f"📡 Testing service at: {BASE_URL}")
    
    # Check if service is running
    try:
        response = requests.get(f"{BASE_URL}/", timeout=5)
        if response.status_code != 200:
            print("❌ Service not running or not accessible")
            return 1
    except Exception as e:
        print(f"❌ Cannot connect to service at {BASE_URL}")
        print(f"   Error: {e}")
        print(f"   Make sure the service is running and accessible")
        return 1
    
    # Create test document
    create_test_docx()
    
    # Test scenarios based on arguments
    if args.light:
        test_scenarios = [
            {"name": "Light Load", "concurrent": 2, "total": 10},
            {"name": "Medium Load", "concurrent": 5, "total": 15},
        ]
    elif args.heavy:
        test_scenarios = [
            {"name": "Heavy Load", "concurrent": 10, "total": 30},
            {"name": "Stress Test", "concurrent": 15, "total": 40},
            {"name": "Extreme Test", "concurrent": 20, "total": 50},
        ]
    else:
        test_scenarios = [
            {"name": "Light Load", "concurrent": 2, "total": 10},
            {"name": "Medium Load", "concurrent": 5, "total": 20},
            {"name": "Heavy Load", "concurrent": 10, "total": 30},
            {"name": "Stress Test", "concurrent": 15, "total": 40},
        ]
    
    results = []
    
    for scenario in test_scenarios:
        print(f"\n🧪 Test Scenario: {scenario['name']}")
        print("-" * 30)
        
        result = run_load_test(
            concurrent_requests=scenario['concurrent'],
            total_requests=scenario['total']
        )
        result['scenario'] = scenario['name']
        results.append(result)
        
        # Wait between tests
        print("   Waiting 10 seconds before next test...")
        time.sleep(10)
    
    # Summary report
    print("\n" + "=" * 50)
    print("📈 PERFORMANCE SUMMARY")
    print("=" * 50)
    
    print(f"{'Scenario':<15} {'Success Rate':<12} {'Avg Time':<10} {'Throughput':<12}")
    print("-" * 50)
    
    for result in results:
        print(f"{result['scenario']:<15} {result['success_rate']:<11.1f}% {result['avg_time']:<9.2f}s {result['throughput']:<11.2f}/s")
    
    # Recommendations
    print(f"\n💡 RECOMMENDATIONS:")
    
    best_throughput = max(results, key=lambda x: x['throughput'])
    best_success = max(results, key=lambda x: x['success_rate'])
    
    print(f"   Best throughput: {best_throughput['scenario']} ({best_throughput['throughput']:.2f}/s)")
    print(f"   Best success rate: {best_success['scenario']} ({best_success['success_rate']:.1f}%)")
    
    # Worker recommendations based on results
    if best_success['success_rate'] < 90:
        print(f"   ⚠️  Consider reducing workers - high failure rate detected")
    elif best_throughput['throughput'] > 2.0:
        print(f"   ✅ Current worker configuration performing well")
    else:
        print(f"   📈 Consider increasing workers if system resources allow")
    
    # Cleanup
    import os
    if os.path.exists('test_worker_doc.docx'):
        os.remove('test_worker_doc.docx')
    
    return 0

if __name__ == "__main__":
    exit(main())
